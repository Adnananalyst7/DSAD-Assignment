from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def style_table(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(10)


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.add_run(text)
    return paragraph


def add_number(doc, text):
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.add_run(text)
    return paragraph


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for index, header in enumerate(headers):
        hdr[index].text = header
        set_cell_shading(hdr[index], "F2F4F7")
        for paragraph in hdr[index].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row_values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            cells[index].text = value
    style_table(table, widths)
    return table


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10

for name, size, color, before, after in [
    ("Heading 1", 16, RGBColor(46, 116, 181), 16, 8),
    ("Heading 2", 13, RGBColor(46, 116, 181), 12, 6),
    ("Heading 3", 12, RGBColor(31, 77, 120), 8, 4),
]:
    style = styles[name]
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.color.rgb = color
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Design Document: Multi-Level Linked List Railway Reservation System")
run.bold = True
run.font.name = "Calibri"
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(11, 37, 69)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.add_run("Group ID: G003 | Assignment 1 - PS2 - Linked List").bold = True

doc.add_heading("1. Problem Overview", level=1)
doc.add_paragraph(
    "The reservation system manages trains, coaches, confirmed passenger reservations, "
    "waiting-list passengers, cancellations, automatic promotion, coach merge checks, "
    "train splitting, hierarchical display, and cycle detection. The implementation uses "
    "a multi-level linked list as the primary structure and reads commands only from "
    "inputPS2.txt while writing results only to outputPS2.txt."
)

doc.add_heading("2. Multi-Level Linked List Design", level=1)
add_bullet(doc, "Level 1 is a singly linked list of TrainNode objects.")
add_bullet(doc, "Each TrainNode stores Train ID, Train Name, Source Station, Destination Station, a pointer to the first CoachNode, and a pointer to the next TrainNode.")
add_bullet(doc, "Level 2 is a singly linked list of CoachNode objects under each train.")
add_bullet(doc, "Each CoachNode stores Coach ID, Coach Type, Maximum Capacity, a confirmed passenger list head, a waiting-list head, and a next coach pointer.")
add_bullet(doc, "Level 3 is a singly linked list of PassengerNode objects for confirmed passengers, plus a separate priority-ordered waiting list for the same coach.")
add_bullet(doc, "No dictionary, set, heap, tree, or external library is used by the submitted Python implementation.")

doc.add_heading("3. Node Fields", level=2)
add_table(
    doc,
    ["Node", "Fields"],
    [
        ["TrainNode", "train_id, train_name, source, destination, coach_head, next"],
        ["CoachNode", "coach_id, coach_type, capacity, passenger_head, waiting_head, next"],
        ["PassengerNode", "passenger_id, name, seat_number, boarding, destination, waiting_priority, is_waiting, next"],
    ],
    [1800, 7560],
)

doc.add_heading("4. Algorithm Explanation", level=1)
doc.add_heading("4.1 Add Train", level=2)
add_number(doc, "Traverse the train linked list to check whether the Train ID already exists.")
add_number(doc, "If not found, create a TrainNode and append it to the end of the train list.")

doc.add_heading("4.2 Add Coach", level=2)
add_number(doc, "Find the required train by linear traversal.")
add_number(doc, "Traverse that train's coach list to prevent duplicate Coach IDs.")
add_number(doc, "Create and append a CoachNode with the specified type and capacity.")

doc.add_heading("4.3 Reserve Ticket", level=2)
add_number(doc, "Find the train and coach through linked-list traversal.")
add_number(doc, "Check both confirmed and waiting passenger lists for duplicate Passenger ID.")
add_number(doc, "If confirmed passenger count is below coach capacity, assign the smallest available seat number and append the passenger to the confirmed list.")
add_number(doc, "If no seat is available, insert the passenger into the waiting list in increasing waiting-priority order. If no explicit priority is provided, arrival order is used.")

doc.add_heading("4.4 Cancel Ticket and Waiting-List Promotion", level=2)
add_number(doc, "Search the confirmed passenger list and unlink the matching passenger if found.")
add_number(doc, "If a waiting-list passenger exists, remove the head of the priority list, assign the freed seat number, and append the passenger to the confirmed list.")
add_number(doc, "If the passenger is not confirmed, search and remove the passenger from the waiting list.")
add_number(doc, "After a confirmed cancellation, adjacent non-empty underutilized coaches are checked for automatic merge.")

doc.add_heading("4.5 Underutilized Coach Merge", level=2)
doc.add_paragraph(
    "Two adjacent coaches are considered merge candidates when both are non-empty, each is below half utilization, "
    "and neither has waiting-list passengers. Confirmed passengers from the second coach are moved to the first coach, "
    "seat numbers are reassigned within the merged coach, capacities are combined, and the second coach is unlinked."
)

doc.add_heading("4.6 Emergency Train Split", level=2)
doc.add_paragraph(
    "For splitTrain::TrainID:CoachID, the system finds the selected coach. If it is not the first coach, the previous "
    "coach's next pointer is set to None and a new TrainNode is appended to the train list. The new train's coach_head "
    "points to the selected coach, preserving the remaining coach chain."
)

doc.add_heading("4.7 Recursive Hierarchical Display", level=2)
doc.add_paragraph(
    "displayTrain and displayAll call recursive functions that print each coach and each passenger list from the current node "
    "to the end of the linked list. Before printing coaches, Floyd's slow-fast pointer cycle detection is executed to avoid "
    "corrupted coach links causing infinite traversal."
)

doc.add_heading("5. Error Handling", level=1)
add_bullet(doc, "Duplicate Train ID, duplicate Coach ID, and duplicate Passenger ID are rejected.")
add_bullet(doc, "Missing trains, coaches, and passengers produce explicit error messages.")
add_bullet(doc, "Invalid command arity and invalid positive integer fields are reported.")
add_bullet(doc, "A non-empty coach cannot be detached.")
add_bullet(doc, "Splitting from the first coach is rejected because it would not create two independent trains.")
add_bullet(doc, "Cycle detection prevents corrupted coach links from being displayed.")

doc.add_heading("6. Runtime Analysis", level=1)
doc.add_paragraph("Let T be the number of trains, C be the number of coaches in a selected train, P be the number of confirmed passengers in a selected coach, W be the number of waiting-list passengers in that coach, and n = T + C + P + W.")
add_table(
    doc,
    ["Operation", "Asymptotic Runtime", "Reason"],
    [
        ["addTrain", "O(T)", "Scans train list for duplicates, then appends."],
        ["removeTrain", "O(T)", "Finds and unlinks one train node."],
        ["addCoach", "O(T + C)", "Finds train and scans coach list for duplicates."],
        ["detachCoach", "O(T + C)", "Finds train, then unlinks an empty coach."],
        ["reserveTicket", "O(T + C + P + W)", "Finds train/coach, checks duplicate passenger, counts passengers, finds seat, or inserts waiting node."],
        ["cancelTicket", "O(T + C + P + W)", "Finds train/coach, unlinks passenger, possibly promotes waiting head and checks adjacent coach merge."],
        ["splitTrain", "O(T + C)", "Finds split point and appends a new train node."],
        ["displayTrain", "O(C + P + W)", "Cycle detection plus recursive traversal of the train hierarchy."],
        ["displayAll", "O(n)", "Visits every train, coach, confirmed passenger, and waiting-list passenger."],
        ["cycle detection", "O(C)", "Floyd slow-fast traversal over coach links."],
    ],
    [1900, 2000, 5460],
)

doc.add_heading("7. File I/O Compliance", level=1)
doc.add_paragraph(
    "The program reads inputPS2.txt from the root folder and writes outputPS2.txt to the same root folder. "
    "No interactive prompts are used and no values are hard coded for evaluation."
)

doc.add_heading("8. Submission Contents", level=1)
add_bullet(doc, "railway_reservation_ps2.py - single Python 3.7 source file.")
add_bullet(doc, "inputPS2.txt - test command file.")
add_bullet(doc, "outputPS2.txt - generated test output.")
add_bullet(doc, "designPS2_G003.docx - design, algorithm, and runtime analysis.")
add_bullet(doc, "G003_Contribution.xlsx - contribution workbook.")

doc.save("designPS2_G003.docx")
